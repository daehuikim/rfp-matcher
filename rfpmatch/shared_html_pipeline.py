from __future__ import annotations

import io
import re
import xml.etree.ElementTree as ET
import zipfile
import unicodedata
import shutil
import subprocess
import tempfile
from html import escape
from copy import copy
from pathlib import Path

import streamlit as st
from bs4 import BeautifulSoup, Tag
try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None

from rfpmatch.shared_html_builder import (
    build_html_document,
    merge_consecutive_tables_in_html,
    merge_consecutive_tables_in_html_raw,
    merge_empty_cells_upward_in_html,
    restore_table_fragments_in_html_raw,
)
from rfpmatch.shared_io import current_artifact_root, save_step1_bundle


def normalize_html_document(html: str) -> str:
    if not html.strip():
        return html
    return str(BeautifulSoup(html, "html.parser"))


def build_txt_from_html_or_markdown(html_text: str, markdown_text: str) -> str:
    if markdown_text.strip():
        return markdown_text
    if html_text.strip():
        soup = BeautifulSoup(html_text, "html.parser")
        return soup.get_text("\n", strip=True)
    return ""


_DOCX_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


def _docx_local_name(tag: str | None) -> str:
    if not tag:
        return ""
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def _docx_element_text(node: ET.Element) -> str:
    parts: list[str] = []
    for elem in node.iter():
        local_name = _docx_local_name(getattr(elem, "tag", None))
        if local_name == "t":
            parts.append(elem.text or "")
        elif local_name == "tab":
            parts.append("\t")
        elif local_name in {"br", "cr"}:
            parts.append("\n")
    return re.sub(r"\s+", " ", "".join(parts)).strip()


def _docx_heading_level(paragraph: ET.Element) -> int | None:
    style = paragraph.find("./w:pPr/w:pStyle", _DOCX_NS)
    style_name = str(style.get(f"{{{_DOCX_NS['w']}}}val", "") if style is not None else "").strip().lower()
    if not style_name:
        return None
    if style_name == "title":
        return 1
    if style_name == "subtitle":
        return 2
    if style_name.startswith("heading"):
        digits = re.sub(r"[^0-9]", "", style_name)
        if digits:
            return max(1, min(int(digits), 6))
        return 1
    return None


def _docx_list_level(paragraph: ET.Element) -> int | None:
    num_pr = paragraph.find("./w:pPr/w:numPr", _DOCX_NS)
    if num_pr is None:
        return None
    ilvl = num_pr.find("./w:ilvl", _DOCX_NS)
    if ilvl is None:
        return 0
    try:
        return max(int(ilvl.get(f"{{{_DOCX_NS['w']}}}val", "0") or 0), 0)
    except (TypeError, ValueError):
        return 0


def _docx_table_rows(table: ET.Element) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.findall("./w:tr", _DOCX_NS):
        cells: list[str] = []
        for cell in row.findall("./w:tc", _DOCX_NS):
            paragraphs = []
            for paragraph in cell.findall("./w:p", _DOCX_NS):
                text = _docx_element_text(paragraph)
                if text:
                    paragraphs.append(text)
            cells.append("<br/>".join(escape(text) for text in paragraphs if text))
        if cells:
            rows.append(cells)
    return rows


def _docx_table_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]
    header = normalized_rows[0]
    separator = ["---"] * width

    def _escape_cell(value: str) -> str:
        return value.replace("|", r"\|").replace("\n", " ").strip()

    lines = [
        "| " + " | ".join(_escape_cell(cell) for cell in header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in normalized_rows[1:]:
        lines.append("| " + " | ".join(_escape_cell(cell) for cell in row) + " |")
    return "\n".join(lines)


def convert_docx_bytes_to_html_markdown_txt(docx_bytes: bytes) -> tuple[str, str, str]:
    if not docx_bytes:
        return "", "", ""

    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zf:
        document_xml = zf.read("word/document.xml")

    root = ET.fromstring(document_xml)
    body = root.find("./w:body", _DOCX_NS)
    if body is None:
        return "", "", ""

    html_parts: list[str] = []
    markdown_parts: list[str] = []
    in_list = False

    def close_list() -> None:
        nonlocal in_list
        if in_list:
            html_parts.append("</ul>")
            in_list = False

    for child in list(body):
        local_name = _docx_local_name(getattr(child, "tag", None))
        if local_name == "p":
            text = _docx_element_text(child)
            if not text:
                continue
            heading_level = _docx_heading_level(child)
            list_level = _docx_list_level(child)
            if heading_level is not None:
                close_list()
                level = max(1, min(heading_level, 6))
                html_parts.append(f"<h{level}>{escape(text)}</h{level}>")
                markdown_parts.append(f"{'#' * level} {text}")
                continue
            if list_level is not None:
                if not in_list:
                    html_parts.append("<ul>")
                    in_list = True
                indent = "  " * min(list_level, 4)
                html_parts.append(f"<li>{escape(text)}</li>")
                markdown_parts.append(f"{indent}- {text}")
                continue
            close_list()
            html_parts.append(f"<p>{escape(text)}</p>")
            markdown_parts.append(text)
            continue
        if local_name == "tbl":
            close_list()
            rows = _docx_table_rows(child)
            if not rows:
                continue
            html_rows = []
            for row_index, row in enumerate(rows):
                cells = "".join(f"<{'th' if row_index == 0 else 'td'}>{cell}</{'th' if row_index == 0 else 'td'}>" for cell in row)
                html_rows.append(f"<tr>{cells}</tr>")
            html_parts.append("<table>" + "".join(html_rows) + "</table>")
            markdown_table = _docx_table_markdown(rows)
            if markdown_table:
                markdown_parts.append(markdown_table)
            continue

    close_list()
    html_body = "".join(html_parts) or "<p></p>"
    html_text = f"<!doctype html><html><head><meta charset='utf-8'></head><body>{html_body}</body></html>"
    markdown_text = "\n\n".join(part for part in markdown_parts if part.strip())
    txt_text = build_txt_from_html_or_markdown(html_text, markdown_text)
    return normalize_html_document(html_text), markdown_text, txt_text


def _find_hwp_converter_command() -> str | None:
    for command in ("hwp5txt", "soffice", "libreoffice"):
        if shutil.which(command):
            return command
    return None


def _text_to_html_document(text: str) -> str:
    body = escape(text or "")
    return normalize_html_document(f"<!doctype html><html><head><meta charset='utf-8'></head><body><pre>{body}</pre></body></html>")


def _write_docx_from_text(text: str, output_path: Path) -> None:
    if Document is None:
        return
    document = Document()
    for line in (text or "").splitlines():
        if line.strip():
            document.add_paragraph(line)
    if not document.paragraphs:
        document.add_paragraph("")
    document.save(output_path)


def _convert_hwp_via_hwp5txt(hwp_path: Path) -> tuple[str, str, str]:
    converter = shutil.which("hwp5txt")
    if not converter:
        raise RuntimeError("hwp5txt 변환 도구를 찾을 수 없습니다.")
    with tempfile.TemporaryDirectory(prefix="rfpmatch_hwp_") as tmpdir:
        tmpdir_path = Path(tmpdir)
        txt_path = tmpdir_path / f"{hwp_path.stem}.txt"
        result = subprocess.run(
            [converter, str(hwp_path), "--output", str(txt_path)],
            capture_output=True,
            text=True,
            check=False,
        )
        if result.returncode != 0:
            raise RuntimeError(
                "hwp5txt를 사용한 HWP 변환에 실패했습니다. "
                f"stderr: {(result.stderr.strip() or result.stdout.strip()) or 'no output'}"
            )
        if not txt_path.exists():
            raise RuntimeError("hwp5txt 변환 결과 TXT 파일을 찾지 못했습니다.")
        txt_text = txt_path.read_text(encoding="utf-8", errors="ignore")
        html_text = _text_to_html_document(txt_text)
        markdown_text = txt_text
        return html_text, markdown_text, txt_text


def detect_document_input_kind(file_name: str, file_bytes: bytes) -> str:
    suffix = Path(str(file_name)).suffix.lower()
    if file_bytes and file_bytes.startswith(b"PK\x03\x04") and zipfile.is_zipfile(io.BytesIO(file_bytes)):
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes), "r") as zf:
                names = {name.lower() for name in zf.namelist()}
        except zipfile.BadZipFile:
            names = set()
        if any(name.startswith("contents/section") and name.endswith(".xml") for name in names):
            return "hwpx"
        if "[content_types].xml" in names and any(name.startswith("word/") for name in names):
            return "docx"
    if file_bytes and file_bytes[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        return "hwp"

    if suffix == ".hwpx":
        return "hwpx"
    if suffix == ".hwp":
        return "hwp"
    if suffix == ".docx":
        return "docx"
    if suffix == ".doc":
        return "doc"
    if suffix == ".pdf":
        return "pdf"
    return "unknown"


def _hwpx_local_name(tag: str | None) -> str:
    if not tag:
        return ""
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def _hwpx_is_noise_text(text: str) -> bool:
    compact = re.sub(r"\s+", " ", (text or "")).strip()
    if not compact:
        return True
    if "원본 그림의 크기" in compact:
        return True
    if compact == "묶음 개체입니다.":
        return True
    return False


def _hwpx_element_text(node: ET.Element) -> str:
    parts: list[str] = []
    for elem in node.iter():
        local_name = _hwpx_local_name(getattr(elem, "tag", None))
        if local_name == "t":
            parts.append(elem.text or "")
        elif local_name == "tab":
            parts.append("\t")
        elif local_name in {"br", "cr"}:
            parts.append("\n")
    text = re.sub(r"\s+", " ", "".join(parts)).strip()
    return "" if _hwpx_is_noise_text(text) else text


def _hwpx_table_rows(table: ET.Element) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.iter():
        if _hwpx_local_name(getattr(row, "tag", None)) != "tr":
            continue
        cells: list[str] = []
        for cell in list(row):
            if _hwpx_local_name(getattr(cell, "tag", None)) != "tc":
                continue
            paragraphs: list[str] = []
            for paragraph in cell.iter():
                if _hwpx_local_name(getattr(paragraph, "tag", None)) == "p":
                    text = _hwpx_element_text(paragraph)
                    if text:
                        paragraphs.append(text)
            cell_text = "<br/>".join(escape(text) for text in paragraphs if text)
            if cell_text:
                cells.append(cell_text)
        if cells:
            rows.append(cells)
    return rows


def _hwpx_sections_from_bytes(hwpx_bytes: bytes) -> list[ET.Element]:
    with zipfile.ZipFile(io.BytesIO(hwpx_bytes), "r") as zf:
        section_names = sorted(
            name for name in zf.namelist() if re.fullmatch(r"Contents/section\d+\.xml", name, flags=re.IGNORECASE)
        )
        sections: list[ET.Element] = []
        for name in section_names:
            try:
                sections.append(ET.fromstring(zf.read(name)))
            except ET.ParseError:
                continue
        return sections


def convert_hwpx_bytes_to_html_markdown_txt(hwpx_bytes: bytes) -> tuple[str, str, str]:
    if not hwpx_bytes:
        return "", "", ""
    sections = _hwpx_sections_from_bytes(hwpx_bytes)
    if not sections:
        raise RuntimeError("HWPX 문서에서 섹션 XML을 찾지 못했습니다.")

    html_parts: list[str] = []
    markdown_parts: list[str] = []

    for section in sections:
        for child in list(section):
            local_name = _hwpx_local_name(getattr(child, "tag", None))
            if local_name == "p":
                text = _hwpx_element_text(child)
                if not text:
                    continue
                html_parts.append(f"<p>{escape(text)}</p>")
                markdown_parts.append(text)
                continue
            if local_name == "tbl":
                rows = _hwpx_table_rows(child)
                if not rows:
                    continue
                html_rows: list[str] = []
                width = max(len(row) for row in rows)
                normalized_rows = [row + [""] * (width - len(row)) for row in rows]
                for row_idx, row in enumerate(normalized_rows):
                    cell_tag = "th" if row_idx == 0 else "td"
                    html_rows.append(
                        "<tr>"
                        + "".join(f"<{cell_tag}>{cell}</{cell_tag}>" for cell in row)
                        + "</tr>"
                    )
                html_parts.append("<table>" + "".join(html_rows) + "</table>")
                markdown_parts.append(
                    "\n".join(
                        "| " + " | ".join(cell.replace("|", r"\|").replace("\n", " ").strip() for cell in row) + " |"
                        for row in normalized_rows
                    )
                )

    html_text = normalize_html_document(f"<!doctype html><html><head><meta charset='utf-8'></head><body>{''.join(html_parts) or '<p></p>'}</body></html>")
    markdown_text = "\n\n".join(part for part in markdown_parts if part.strip())
    txt_text = build_txt_from_html_or_markdown(html_text, markdown_text)
    return html_text, markdown_text, txt_text


def convert_hwp_file_to_html_markdown_txt(hwp_path: Path) -> tuple[str, str, str]:
    if not hwp_path.exists():
        raise FileNotFoundError(hwp_path)

    converter = _find_hwp_converter_command()
    if not converter:
        raise RuntimeError(
            "HWP 파일을 변환할 수 있는 도구가 없습니다. "
            "현재 환경에는 hwp5txt나 LibreOffice/soffice가 없어 HWP를 직접 HTML로 변환할 수 없습니다."
        )

    if converter == "hwp5txt":
        return _convert_hwp_via_hwp5txt(hwp_path)

    with tempfile.TemporaryDirectory(prefix="rfpmatch_hwp_") as tmpdir:
        tmpdir_path = Path(tmpdir)
        attempts: list[str] = []
        for convert_format in ("html", "docx", "txt"):
            result = subprocess.run(
                [
                    converter,
                    "--headless",
                    "--convert-to",
                    convert_format,
                    "--outdir",
                    str(tmpdir_path),
                    str(hwp_path),
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            stderr_text = (result.stderr.strip() or result.stdout.strip()).strip()
            if result.returncode != 0:
                attempts.append(f"{convert_format}: {stderr_text or 'conversion failed'}")
                continue

            if convert_format == "html":
                converted_html = tmpdir_path / f"{hwp_path.stem}.html"
                if not converted_html.exists():
                    html_candidates = list(tmpdir_path.glob("*.html"))
                    if html_candidates:
                        converted_html = html_candidates[0]
                if converted_html.exists():
                    html_text = normalize_html_document(converted_html.read_text(encoding="utf-8", errors="ignore"))
                    markdown_text = BeautifulSoup(html_text, "html.parser").get_text("\n", strip=True)
                    txt_text = build_txt_from_html_or_markdown(html_text, markdown_text)
                    return html_text, markdown_text, txt_text
                attempts.append(f"{convert_format}: output file missing")
                continue

            if convert_format == "docx":
                converted_docx = tmpdir_path / f"{hwp_path.stem}.docx"
                if not converted_docx.exists():
                    docx_candidates = list(tmpdir_path.glob("*.docx"))
                    if docx_candidates:
                        converted_docx = docx_candidates[0]
                if converted_docx.exists():
                    return convert_docx_bytes_to_html_markdown_txt(converted_docx.read_bytes())
                attempts.append(f"{convert_format}: output file missing")
                continue

            converted_txt = tmpdir_path / f"{hwp_path.stem}.txt"
            if not converted_txt.exists():
                txt_candidates = list(tmpdir_path.glob("*.txt"))
                if txt_candidates:
                    converted_txt = txt_candidates[0]
            if converted_txt.exists():
                txt_text = converted_txt.read_text(encoding="utf-8", errors="ignore")
                html_text = normalize_html_document(f"<html><head><meta charset='utf-8'></head><body><pre>{escape(txt_text)}</pre></body></html>")
                markdown_text = txt_text
                return html_text, markdown_text, txt_text
            attempts.append(f"{convert_format}: output file missing")

        raise RuntimeError(
            "HWP 변환 결과 HTML/DOCX/TXT 파일을 찾지 못했습니다. "
            f"시도 결과: {'; '.join(attempts) if attempts else 'no converter output'}"
        )


def convert_hwpx_file_to_html_markdown_txt(hwpx_path: Path) -> tuple[str, str, str]:
    if not hwpx_path.exists():
        raise FileNotFoundError(hwpx_path)
    return convert_hwpx_bytes_to_html_markdown_txt(hwpx_path.read_bytes())


def write_docx_step1_outputs(docx_path: Path, output_root: Path) -> dict[str, Path]:
    html_text, markdown_text, txt_text = convert_docx_bytes_to_html_markdown_txt(docx_path.read_bytes())
    output_root.mkdir(parents=True, exist_ok=True)
    html_path = output_root / f"{docx_path.stem}.html"
    markdown_path = output_root / f"{docx_path.stem}.md"
    txt_path = output_root / f"{docx_path.stem}.txt"
    html_path.write_text(html_text, encoding="utf-8")
    markdown_path.write_text(markdown_text, encoding="utf-8")
    txt_path.write_text(txt_text, encoding="utf-8")
    return {"html": html_path, "markdown": markdown_path, "txt": txt_path}


def write_hwp_step1_outputs(hwp_path: Path, output_root: Path) -> dict[str, Path]:
    html_text, markdown_text, txt_text = convert_hwp_file_to_html_markdown_txt(hwp_path)
    output_root.mkdir(parents=True, exist_ok=True)
    html_path = output_root / f"{hwp_path.stem}.html"
    markdown_path = output_root / f"{hwp_path.stem}.md"
    txt_path = output_root / f"{hwp_path.stem}.txt"
    docx_path = output_root / f"{hwp_path.stem}.docx"
    html_path.write_text(html_text, encoding="utf-8")
    markdown_path.write_text(markdown_text, encoding="utf-8")
    txt_path.write_text(txt_text, encoding="utf-8")
    _write_docx_from_text(txt_text, docx_path)
    outputs = {"html": html_path, "markdown": markdown_path, "txt": txt_path}
    if docx_path.exists():
        outputs["docx"] = docx_path
    return outputs


def write_hwpx_step1_outputs(hwpx_path: Path, output_root: Path) -> dict[str, Path]:
    html_text, markdown_text, txt_text = convert_hwpx_file_to_html_markdown_txt(hwpx_path)
    output_root.mkdir(parents=True, exist_ok=True)
    html_path = output_root / f"{hwpx_path.stem}.html"
    markdown_path = output_root / f"{hwpx_path.stem}.md"
    txt_path = output_root / f"{hwpx_path.stem}.txt"
    html_path.write_text(html_text, encoding="utf-8")
    markdown_path.write_text(markdown_text, encoding="utf-8")
    txt_path.write_text(txt_text, encoding="utf-8")
    return {"html": html_path, "markdown": markdown_path, "txt": txt_path}


def strip_html_styles(html: str) -> str:
    if not html.strip():
        return html
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup.find_all(["style", "script"]):
        tag.decompose()
    for tag in soup.find_all("span"):
        tag.unwrap()
    for tag in list(soup.find_all(True)):
        if tag.has_attr("style"):
            del tag["style"]
    return str(soup)


def clean_inline_text(value: str) -> str:
    return " ".join((value or "").split()).strip()


def is_effectively_empty_cell(cell: Tag | None) -> bool:
    if cell is None:
        return True
    text = clean_inline_text(cell.get_text(" ", strip=False))
    return not bool(text)


def prune_all_empty_table_columns_in_html(html: str) -> str:
    if not html.strip():
        return html
    current_html = html
    for _ in range(4):
        soup = BeautifulSoup(current_html, "html.parser")
        changed = False
        last_table_col_count = None
        for table in soup.find_all("table"):
            tr_nodes = [tr for tr in table.find_all("tr") if tr.find_parent("table") is table]
            if not tr_nodes:
                continue
            grid: dict[tuple[int, int], Tag] = {}
            cell_pos: dict[int, tuple[int, int]] = {}
            max_col = 0
            for r_idx, tr in enumerate(tr_nodes, start=0):
                c_idx = 0
                while (r_idx, c_idx) in grid:
                    c_idx += 1
                for cell in tr.find_all(["th", "td"], recursive=False):
                    while (r_idx, c_idx) in grid:
                        c_idx += 1
                    try:
                        rowspan = max(int(cell.get("rowspan", 1) or 1), 1)
                    except (TypeError, ValueError):
                        rowspan = 1
                    try:
                        colspan = max(int(cell.get("colspan", 1) or 1), 1)
                    except (TypeError, ValueError):
                        colspan = 1
                    cell_pos[id(cell)] = (r_idx, c_idx)
                    for dr in range(rowspan):
                        for dc in range(colspan):
                            grid[(r_idx + dr, c_idx + dc)] = cell
                    c_idx += colspan
                    max_col = max(max_col, c_idx)
            if max_col <= 1:
                if max_col == 1:
                    last_table_col_count = 1
                continue

            def visible_value_count(col_idx: int) -> int:
                visible = 0
                for r_idx in range(len(tr_nodes)):
                    cell = grid.get((r_idx, col_idx))
                    if cell is None:
                        continue
                    if clean_inline_text(cell.get_text(" ", strip=False)):
                        visible += 1
                return visible

            active_cols = [c for c in range(max_col) if visible_value_count(c) > 0]
            removable_cols: set[int] = set()
            
            if active_cols:
                active_cols_count = len(active_cols)
                left_empty_count = min(active_cols)
                
                if last_table_col_count is not None and last_table_col_count > 0:
                    needed_left_empty = max(0, last_table_col_count - active_cols_count)
                    allowed_to_remove = max(0, left_empty_count - needed_left_empty)
                    if allowed_to_remove > 0 and visible_value_count(0) == 0:
                        removable_cols.add(0)
                else:
                    if visible_value_count(0) == 0:
                        removable_cols.add(0)
                
                if max_col - 1 != 0 and visible_value_count(max_col - 1) == 0:
                    removable_cols.add(max_col - 1)
            else:
                if visible_value_count(0) == 0:
                    removable_cols.add(0)
                if max_col - 1 != 0 and visible_value_count(max_col - 1) == 0:
                    removable_cols.add(max_col - 1)

            if not removable_cols:
                last_table_col_count = max_col
                continue
            new_rows: list[Tag] = []
            for r_idx, tr in enumerate(tr_nodes, start=0):
                row_tag = soup.new_tag("tr")
                for cell in tr.find_all(["th", "td"], recursive=False):
                    start_col = cell_pos.get(id(cell), (r_idx, 0))[1]
                    try:
                        colspan = max(int(cell.get("colspan", 1) or 1), 1)
                    except (TypeError, ValueError):
                        colspan = 1
                    cell_cols = set(range(start_col, start_col + colspan))
                    if cell_cols & removable_cols:
                        continue
                    row_tag.append(copy(cell))
                if row_tag.find_all(["th", "td"], recursive=False):
                    new_rows.append(row_tag)
            if new_rows:
                table.clear()
                for row_tag in new_rows:
                    table.append(row_tag)
                changed = True
            
            remaining_col_count = max_col - len(removable_cols)
            last_table_col_count = max(1, remaining_col_count)
        next_html = str(soup) if changed else current_html
        if next_html == current_html:
            return current_html
        current_html = next_html
    return current_html


def selected_pdf_basename() -> str:
    selected = str(st.session_state.get("uploaded_pdf_name") or st.session_state.get("file_name") or "").strip()
    return unicodedata.normalize("NFC", Path(selected).name)


def is_hana_pdf_selected() -> bool:
    target_names = ("하나.pdf", "하나은행.pdf")
    candidates = [
        st.session_state.get("uploaded_pdf_path"),
        st.session_state.get("uploaded_pdf_name"),
        st.session_state.get("file_name"),
    ]
    for candidate in candidates:
        normalized = unicodedata.normalize("NFC", str(candidate or "")).strip()
        if normalized and any(target_name in normalized for target_name in target_names):
            return True
    return False


def apply_empty_table_column_postprocess_to_html(raw_html: str) -> str:
    if not raw_html.strip():
        return raw_html
    return prune_all_empty_table_columns_in_html(raw_html)


def apply_shinhanlife_table_exception_in_raw_html(raw_html: str, file_name: str | None) -> str:
    if not raw_html.strip():
        return raw_html
    normalized_name = unicodedata.normalize("NFC", str(file_name or ""))
    if "신한라이프" not in normalized_name:
        return raw_html
    soup = BeautifulSoup(raw_html, "html.parser")
    changed = False
    target_labels = ["물리 구성도", "H/W S/W Spec", "Storage", "백업", "호환성"]
    for table in soup.find_all("table"):
        tr_nodes = [tr for tr in table.find_all("tr") if tr.find_parent("table") is table]
        if len(tr_nodes) != 6:
            continue
        first_row_cells = tr_nodes[0].find_all(["th", "td"], recursive=False)
        if len(first_row_cells) != 2:
            continue
        if clean_inline_text(first_row_cells[1].get_text(" ", strip=True)) != "- Pod 및 Container 간 연계를 포함하는 논리 구성도 제시":
            continue
        matched = True
        for tr, expected_label in zip(tr_nodes[1:], target_labels):
            cells = tr.find_all(["th", "td"], recursive=False)
            if len(cells) != 2 or clean_inline_text(cells[0].get_text(" ", strip=True)) != expected_label:
                matched = False
                break
        if not matched:
            continue
        lead_cell = first_row_cells[0]
        if not is_effectively_empty_cell(lead_cell):
            continue
        lead_cell["rowspan"] = str(len(tr_nodes))
        changed = True
        break
    return str(soup) if changed else raw_html


def build_step1_raw_html_variants(
    raw_html_original: str,
    *,
    file_name: str | None = None,
    apply_empty_column_cleanup: bool = False,
) -> dict[str, str]:
    """Build the step1 raw HTML variants used for debugging and downstream steps."""
    original_html = raw_html_original or ""
    restored_html = original_html
    if apply_empty_column_cleanup:
        restored_html = apply_empty_table_column_postprocess_to_html(restored_html)
    merged_html = merge_consecutive_tables_in_html_raw(restored_html)
    merged_up_html = merge_empty_cells_upward_in_html(merged_html)
    return {
        "html_raw_original": original_html,
        "html_raw_restored": restored_html,
        "html_raw": restored_html,
        "html_raw_merged_empty_up_postprocessed": merged_up_html,
        "html_merged_from_raw": merged_html,
        "html_merged_from_raw_postprocessed": merged_up_html,
    }


def text_from_content_block(block: dict) -> str:
    texts: list[str] = []
    for key in ("content", "text", "title", "caption"):
        value = block.get(key)
        if isinstance(value, str) and value.strip():
            texts.append(value.strip())
    for line in block.get("lines", []):
        for span in line.get("spans", []):
            content = span.get("content")
            if isinstance(content, str) and content.strip():
                texts.append(content.strip())
    for key in ("items", "children", "blocks", "elements"):
        value = block.get(key)
        if isinstance(value, list):
            for item in value:
                if isinstance(item, dict):
                    for inner_key in ("content", "text", "title"):
                        inner_val = item.get(inner_key)
                        if isinstance(inner_val, str) and inner_val.strip():
                            texts.append(inner_val.strip())
    return " ".join(texts).strip()


def build_page_summary(content_list: list[dict]) -> list[dict]:
    pages: dict[int, dict] = {}
    for block in content_list:
        page_idx = block.get("page_idx")
        if not isinstance(page_idx, int):
            page_raw = block.get("page_no", block.get("page"))
            if isinstance(page_raw, int):
                page_idx = page_raw if page_raw > 0 else page_raw
        if not isinstance(page_idx, int):
            continue
        info = pages.setdefault(page_idx, {"page": page_idx, "block_count": 0, "sample": ""})
        info["block_count"] += 1
        if not info["sample"]:
            sample = text_from_content_block(block)
            if sample:
                info["sample"] = sample[:140]
    return [pages[key] for key in sorted(pages.keys())]


def build_page_summary_from_html(html_text: str) -> list[dict]:
    soup = BeautifulSoup(html_text or "", "html.parser")
    body = soup.body or soup
    pages: dict[int, dict] = {}
    direct_blocks = [child for child in body.find_all(recursive=False) if getattr(child, "attrs", None)]
    if not direct_blocks:
        direct_blocks = body.find_all(attrs={"data-page": True})

    for block in direct_blocks:
        page_raw = block.get("data-page")
        page_idx: int | None = None
        if isinstance(page_raw, int):
            page_idx = page_raw
        elif isinstance(page_raw, str):
            match = re.search(r"\d+", page_raw)
            if match:
                try:
                    page_idx = int(match.group(0))
                except ValueError:
                    page_idx = None
        if page_idx is None:
            continue
        if page_idx < 0:
            continue
        info = pages.setdefault(page_idx, {"page": page_idx, "block_count": 0, "sample": ""})
        info["block_count"] += 1
        if not info["sample"]:
            sample = re.sub(r"\s+", " ", block.get_text(" ", strip=True)).strip()
            if sample:
                info["sample"] = sample[:140]
    return [pages[key] for key in sorted(pages.keys())]


def html_table_merge_summary(html_text: str) -> dict[str, int]:
    soup = BeautifulSoup(html_text or "", "html.parser")
    tables = soup.find_all("table")
    rowspan_cells = 0
    for table in tables:
        for cell in table.find_all(["td", "th"]):
            try:
                rowspan = int(cell.get("rowspan", 1) or 1)
            except (TypeError, ValueError):
                rowspan = 1
            if rowspan > 1:
                rowspan_cells += 1
    return {"tables": len(tables), "rowspan_cells": rowspan_cells}


def merge_current_html_table_sections(html_text: str, *, source_key: str, source_label: str, output_suffix: str) -> bool:
    if not html_text.strip():
        return False
    before_summary = html_table_merge_summary(html_text)
    merged_html = merge_consecutive_tables_in_html_raw(html_text) if source_key == "raw" else merge_consecutive_tables_in_html(html_text)
    if merged_html.strip() == html_text.strip():
        return False
    after_summary = html_table_merge_summary(merged_html)
    summary = {
        "before_tables": before_summary["tables"],
        "after_tables": after_summary["tables"],
        "merged_tables": max(before_summary["tables"] - after_summary["tables"], 0),
        "rowspan_cells": after_summary["rowspan_cells"],
    }
    st.session_state["html_merged"] = merged_html
    st.session_state["html_merged_raw"] = merged_html
    st.session_state["html_merged_source"] = source_label
    if source_key == "raw":
        st.session_state["html_merged_from_raw"] = merged_html
        st.session_state["html_merged_summary_from_raw"] = summary
    elif source_key == "postprocessed":
        st.session_state["html_merged_from_postprocessed"] = merged_html
        st.session_state["html_merged_summary_from_postprocessed"] = summary
    st.session_state["html_merged_summary"] = summary
    artifact_root = current_artifact_root()
    artifact_root.mkdir(parents=True, exist_ok=True)
    merged_path = artifact_root / f"document_merged_tables_{output_suffix}.html"
    try:
        merged_path.write_text(merged_html, encoding="utf-8")
        outputs = dict(st.session_state.get("outputs") or {})
        outputs.pop("html_merged", None)
        outputs[f"html_merged_from_{source_key}"] = str(merged_path.resolve())
        st.session_state["outputs"] = outputs
    except OSError:
        pass
    bundle_path = save_step1_bundle(include_merge_outputs=True)
    if bundle_path:
        st.success(f"{source_label} 기준 표 병합 후 1단계 복원 번들을 저장했습니다: {bundle_path}")
    return True


def postprocess_merged_html_from_raw(*, source_label: str = "HTML 문서 연속표 병합본") -> bool:
    merged_html = st.session_state.get("html_merged_from_raw") or ""
    if not merged_html.strip():
        return False
    postprocessed_html = merge_empty_cells_upward_in_html(merged_html)
    before_summary = html_table_merge_summary(merged_html)
    after_summary = html_table_merge_summary(postprocessed_html)
    summary = {
        "before_tables": before_summary["tables"],
        "after_tables": after_summary["tables"],
        "merged_tables": max(before_summary["tables"] - after_summary["tables"], 0),
        "rowspan_cells": after_summary["rowspan_cells"],
    }
    st.session_state["html_raw_merged_empty_up_postprocessed"] = postprocessed_html
    st.session_state["html_merged_from_raw_postprocessed"] = postprocessed_html
    st.session_state["html_merged_summary_from_raw_postprocessed"] = summary
    st.session_state["html_merged"] = postprocessed_html
    st.session_state["html_merged_summary"] = summary
    artifact_root = current_artifact_root()
    artifact_root.mkdir(parents=True, exist_ok=True)
    postprocessed_path = artifact_root / "document_html_raw_단순병합_빈칸윗칸병합_postprocessed.html"
    try:
        postprocessed_path.write_text(postprocessed_html, encoding="utf-8")
        outputs = dict(st.session_state.get("outputs") or {})
        outputs["html_raw_merged_empty_up_postprocessed"] = str(postprocessed_path.resolve())
        outputs["html_merged_from_raw_postprocessed"] = str(postprocessed_path.resolve())
        st.session_state["outputs"] = outputs
    except OSError:
        pass
    bundle_path = save_step1_bundle(include_merge_outputs=True)
    if bundle_path:
        st.success(f"{source_label} 후처리본을 저장했습니다: {bundle_path}")
    return True


def fallback_html_from_markdown(outputs: dict, content_list: list[dict]) -> str:
    return normalize_html_document(build_html_document(content_list, outputs["markdown"]))
